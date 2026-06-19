import os
import logging
import zipfile
import tempfile
import xml.etree.ElementTree as ET
import copy
import shutil
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

logger = logging.getLogger("cis-docx-engine")

# XML namespaces for DOCX format
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
}
W_NAMESPACE = NAMESPACES['w']
W = f"{{{W_NAMESPACE}}}"

# Register namespaces in ElementTree to preserve prefixes on save
for prefix, uri in NAMESPACES.items():
    ET.register_namespace(prefix, uri)

def get_p_text(p_el) -> str:
    """Helper to extract text content from an ElementTree paragraph node."""
    text_parts = []
    for t in p_el.findall(f".//{W}t"):
        if t.text:
            text_parts.append(t.text)
    return "".join(text_parts)

class DocxEngine:
    def merge_changes(self, template_path: str, optimized_sections: dict, output_path: str, placeholders: dict = None):
        """
        Merges optimized sections back into a DOCX layout.
        Uses direct XML manipulation to guarantee 100% style and visual element preservation.
        Falls back to python-docx text replacement if direct XML fails.
        """
        logger.info(f"Merging changes into template {template_path}")
        
        # Ensure the template exists; if not, create a base styled template
        if not os.path.exists(template_path):
            self._create_default_template(template_path)
            
        try:
            # Attempt direct XML/ZIP manipulation
            success = self.unzip_and_replace_xml(template_path, optimized_sections, output_path, placeholders)
            if success:
                logger.info("Direct XML/ZIP manipulation succeeded.")
                return
        except Exception as e:
            logger.error(f"Direct XML AST manipulation failed: {e}. Falling back to python-docx...")

        # Fallback to traditional python-docx manipulation
        self._merge_changes_docx_fallback(template_path, optimized_sections, output_path)

    def unzip_and_replace_xml(self, template_path: str, optimized_sections: dict, output_path: str, placeholders: dict = None) -> bool:
        """
        Directly unzips the DOCX package, parses word/document.xml using ElementTree,
        performs placeholder interpolation and styling-preserving replacements,
        runs layout structure regression tests, and writes the zip back.
        """
        temp_dir = tempfile.mkdtemp()
        try:
            # Unzip template files
            with zipfile.ZipFile(template_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
                
            doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
            if not os.path.exists(doc_xml_path):
                logger.error("word/document.xml not found in template package.")
                return False
                
            with open(doc_xml_path, 'r', encoding='utf-8') as f:
                original_xml = f.read()
                
            root = ET.fromstring(original_xml)
            
            # 1. Perform placeholder interpolation
            if placeholders:
                for t_el in root.findall(f".//{W}t"):
                    if t_el.text:
                        for k, v in placeholders.items():
                            ph = f"{{{{{k}}}}}"
                            if ph in t_el.text:
                                t_el.text = t_el.text.replace(ph, v)
                                
            # 2. Perform Content Control (w:sdt) replacements
            for sdt in root.findall(f".//{W}sdt"):
                sdtPr = sdt.find(f"{W}sdtPr")
                if sdtPr is not None:
                    alias = sdtPr.find(f"{W}alias")
                    alias_val = alias.get(f"{W}val") if alias is not None else None
                    if alias_val and alias_val in optimized_sections:
                        sdtContent = sdt.find(f"{W}sdtContent")
                        if sdtContent is not None:
                            bullets = [b.strip() for b in optimized_sections[alias_val].split("\n") if b.strip()]
                            p_list = sdtContent.findall(f"{W}p")
                            
                            for idx, b_text in enumerate(bullets):
                                if idx < len(p_list):
                                    self._replace_xml_p_text(p_list[idx], b_text)
                                else:
                                    if p_list:
                                        new_p = copy.deepcopy(p_list[-1])
                                        self._replace_xml_p_text(new_p, b_text)
                                        sdtContent.append(new_p)
                                        p_list.append(new_p)
                                        
                            if len(p_list) > len(bullets):
                                for p_to_clear in p_list[len(bullets):]:
                                    self._replace_xml_p_text(p_to_clear, "")
                                    
            # 3. Perform Heading-based section replacements (standard body paragraphs)
            body = root.find(f"{W}body")
            if body is not None:
                body_children = list(body)
                for section_name, optimized_content in optimized_sections.items():
                    if not optimized_content:
                        continue
                        
                    # Find heading paragraph matching the section title
                    heading_idx = -1
                    for idx, el in enumerate(body_children):
                        if el.tag == f"{W}p" and get_p_text(el).strip().lower() == section_name.lower():
                            heading_idx = idx
                            break
                            
                    if heading_idx == -1:
                        continue
                        
                    if section_name.lower() == "professional experience":
                        opt_bullets = [
                            b.strip() for b in optimized_content.split("\n") if b.strip()
                        ]
                        
                        # Find existing bullet paragraphs below this heading
                        bullet_elements = []
                        idx = heading_idx + 1
                        while idx < len(body_children):
                            el = body_children[idx]
                            if el.tag == f"{W}p":
                                p_text = get_p_text(el).strip()
                                pPr = el.find(f"{W}pPr")
                                style = ""
                                if pPr is not None:
                                    pStyle = pPr.find(f"{W}pStyle")
                                    if pStyle is not None:
                                        style = pStyle.get(f"{W}val", "")
                                        
                                is_heading = style.startswith("Heading") or p_text in ["Skills Summary", "Education", "Certifications"]
                                if is_heading:
                                    break
                                    
                                if style.startswith("List") or p_text.startswith("-") or p_text.startswith("*") or p_text.startswith("•"):
                                    bullet_elements.append(el)
                                elif p_text == "":
                                    pass
                                else:
                                    if bullet_elements:
                                        break
                            idx += 1
                            
                        # Replace in-place or append
                        for i, opt_text in enumerate(opt_bullets):
                            if i < len(bullet_elements):
                                self._replace_xml_p_text(bullet_elements[i], opt_text)
                            else:
                                if bullet_elements:
                                    last_el = bullet_elements[-1]
                                    new_el = copy.deepcopy(last_el)
                                    self._replace_xml_p_text(new_el, opt_text)
                                    
                                    body_el_list = list(body)
                                    tree_idx = body_el_list.index(last_el)
                                    body.insert(tree_idx + 1, new_el)
                                    bullet_elements.append(new_el)
                                else:
                                    p = ET.SubElement(body, f"{W}p")
                                    self._replace_xml_p_text(p, opt_text)
                                    
                        # Clear extra paragraphs
                        if len(bullet_elements) > len(opt_bullets):
                            for bullet_el in bullet_elements[len(opt_bullets):]:
                                self._replace_xml_p_text(bullet_el, "")
                                
                    else:
                        # Single-paragraph section replacement (Skills Summary, Education, etc.)
                        idx = heading_idx + 1
                        target_p = None
                        while idx < len(body_children):
                            el = body_children[idx]
                            if el.tag == f"{W}p":
                                p_text = get_p_text(el).strip()
                                pPr = el.find(f"{W}pPr")
                                style = ""
                                if pPr is not None:
                                    pStyle = pPr.find(f"{W}pStyle")
                                    if pStyle is not None:
                                        style = pStyle.get(f"{W}val", "")
                                if style.startswith("Heading"):
                                    break
                                if p_text != "":
                                    target_p = el
                                    break
                            idx += 1
                            
                        if target_p is not None:
                            self._replace_xml_p_text(target_p, optimized_content.strip())
                            
            # Convert tree back to XML string
            modified_xml = ET.tostring(root, encoding='utf-8').decode('utf-8')
            
            # Layout structural regression validation
            if not self.verify_layout_regression(original_xml, modified_xml):
                raise ValueError("Visual Layout Preservation Check Failed: Structural elements (margins/tables) mutated.")
                
            # Write modified XML back
            with open(doc_xml_path, 'w', encoding='utf-8') as f:
                f.write(modified_xml)
                
            # Repackage files into output zip/docx
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                for root_dir, _, files in os.walk(temp_dir):
                    for file in files:
                        full_path = os.path.join(root_dir, file)
                        rel_path = os.path.relpath(full_path, temp_dir)
                        zip_out.write(full_path, rel_path)
                        
            logger.info(f"Direct XML AST manipulation completed successfully. Saved to {output_path}")
            return True
            
        finally:
            shutil.rmtree(temp_dir)

    def _replace_xml_p_text(self, p_el, new_text: str):
        """Modifies paragraph runs at the ElementTree XML level to preserve run format styles."""
        # Clean bullet characters
        new_text = new_text.lstrip("-*• ").strip()
        
        runs = p_el.findall(f".//{W}r")
        if not runs:
            r = ET.SubElement(p_el, f"{W}r")
            t = ET.SubElement(r, f"{W}t")
            t.text = new_text
            return
            
        first_t = None
        for r in runs:
            t_el = r.find(f"{W}t")
            if t_el is not None:
                if first_t is None:
                    first_t = t_el
                    first_t.text = new_text
                else:
                    t_el.text = ""
                    
        if first_t is None:
            first_t = ET.SubElement(runs[0], f"{W}t")
            first_t.text = new_text

    def verify_layout_regression(self, original_xml: str, modified_xml: str) -> bool:
        """
        Validates layout preservation. Ensures tables count, page settings, 
        and margins are structurally identical before and after.
        """
        orig_tree = ET.fromstring(original_xml)
        mod_tree = ET.fromstring(modified_xml)
        
        # 1. Check table count
        orig_tbls = len(orig_tree.findall(f".//{W}tbl"))
        mod_tbls = len(mod_tree.findall(f".//{W}tbl"))
        if orig_tbls != mod_tbls:
            logger.error(f"Regression: Table count changed from {orig_tbls} to {mod_tbls}")
            return False
            
        # 2. Check margins and page settings
        orig_pgMar = orig_tree.find(f".//{W}pgMar")
        mod_pgMar = mod_tree.find(f".//{W}pgMar")
        if orig_pgMar is not None and mod_pgMar is not None:
            for attr in ['top', 'bottom', 'left', 'right', 'header', 'footer']:
                # Handle attributes with or without namespace w prefix in find
                orig_val = orig_pgMar.get(f"{W}{attr}") or orig_pgMar.get(attr)
                mod_val = mod_pgMar.get(f"{W}{attr}") or mod_pgMar.get(attr)
                if orig_val != mod_val:
                    logger.error(f"Regression: Page margin '{attr}' changed from {orig_val} to {mod_val}")
                    return False
                    
        orig_pgSz = orig_tree.find(f".//{W}pgSz")
        mod_pgSz = mod_tree.find(f".//{W}pgSz")
        if orig_pgSz is not None and mod_pgSz is not None:
            for attr in ['w', 'h', 'orient']:
                orig_val = orig_pgSz.get(f"{W}{attr}") or orig_pgSz.get(attr)
                mod_val = mod_pgSz.get(f"{W}{attr}") or mod_pgSz.get(attr)
                if orig_val != mod_val:
                    logger.error(f"Regression: Page size/orientation '{attr}' changed from {orig_val} to {mod_val}")
                    return False
                    
        return True

    def _merge_changes_docx_fallback(self, template_path: str, optimized_sections: dict, output_path: str):
        """Traditional python-docx layout replacement fallback."""
        doc = Document(template_path)
        
        for section_name, optimized_content in optimized_sections.items():
            if not optimized_content:
                continue
                
            heading_idx = -1
            for idx, p in enumerate(doc.paragraphs):
                if p.text.strip().lower() == section_name.lower():
                    heading_idx = idx
                    break
                    
            if heading_idx == -1:
                continue
                
            if section_name.lower() == "professional experience":
                opt_bullets = [
                    b.strip().lstrip("-*• ").strip()
                    for b in optimized_content.split("\n")
                    if b.strip()
                ]
                
                bullet_paragraphs = []
                idx = heading_idx + 1
                while idx < len(doc.paragraphs):
                    p = doc.paragraphs[idx]
                    p_text = p.text.strip()
                    
                    if p.style.name.startswith("Heading") or (p_text and p_text in ["Skills Summary", "Education", "Certifications"]):
                        break
                        
                    if p.style.name.startswith("List") or p_text.startswith("-") or p_text.startswith("*") or p_text.startswith("•"):
                        bullet_paragraphs.append(p)
                    elif p_text == "":
                        pass
                    else:
                        if bullet_paragraphs:
                            break
                    idx += 1
                
                for i, opt_text in enumerate(opt_bullets):
                    if i < len(bullet_paragraphs):
                        self._replace_paragraph_text_preserving_style(bullet_paragraphs[i], opt_text)
                    else:
                        if bullet_paragraphs:
                            last_bullet = bullet_paragraphs[-1]
                            new_p = self._insert_paragraph_after(last_bullet, opt_text, style=last_bullet.style)
                            self._copy_paragraph_format(last_bullet, new_p)
                            bullet_paragraphs.append(new_p)
                        else:
                            p = doc.add_paragraph(opt_text, style='List Bullet')
                            p.paragraph_format.left_indent = Inches(0.5)
                            bullet_paragraphs.append(p)
                            
                if len(bullet_paragraphs) > len(opt_bullets):
                    for excess_p in bullet_paragraphs[len(opt_bullets):]:
                        self._replace_paragraph_text_preserving_style(excess_p, "")
                        
            else:
                idx = heading_idx + 1
                target_p = None
                while idx < len(doc.paragraphs):
                    p = doc.paragraphs[idx]
                    if p.style.name.startswith("Heading"):
                        break
                    if p.text.strip() != "":
                        target_p = p
                        break
                    idx += 1
                    
                if target_p:
                    self._replace_paragraph_text_preserving_style(target_p, optimized_content.strip())
                    
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

    def _replace_paragraph_text_preserving_style(self, paragraph: Paragraph, new_text: str):
        new_text = new_text.lstrip("-*• ").strip()
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""

    def _insert_paragraph_after(self, paragraph: Paragraph, text: str, style=None) -> Paragraph:
        p_element = paragraph._p
        parent = p_element.getparent()
        new_p_element = OxmlElement('w:p')
        p_element.addnext(new_p_element)
        new_para = Paragraph(new_p_element, paragraph._parent)
        if text:
            new_para.text = text
        if style:
            new_para.style = style
        return new_para

    def _copy_paragraph_format(self, src: Paragraph, dst: Paragraph):
        dst.paragraph_format.left_indent = src.paragraph_format.left_indent
        dst.paragraph_format.right_indent = src.paragraph_format.right_indent
        dst.paragraph_format.space_before = src.paragraph_format.space_before
        dst.paragraph_format.space_after = src.paragraph_format.space_after
        dst.paragraph_format.line_spacing = src.paragraph_format.line_spacing
        
        if src.runs and dst.runs:
            dst_run = dst.runs[0]
            src_run = src.runs[0]
            dst_run.bold = src_run.bold
            dst_run.italic = src_run.italic
            if src_run.font.name:
                dst_run.font.name = src_run.font.name
            if src_run.font.size:
                dst_run.font.size = src_run.font.size
            if src_run.font.color and src_run.font.color.rgb:
                dst_run.font.color.rgb = src_run.font.color.rgb

    def _create_default_template(self, path: str):
        os.makedirs(os.path.dirname(path), exist_ok=True)
        doc = Document()
        
        title = doc.add_paragraph()
        run = title.add_run("John Doe")
        run.font.size = Pt(24)
        run.bold = True
        
        contact = doc.add_paragraph("Email: john.doe@example.com | Phone: 123-456-7890 | San Francisco, CA")
        contact.paragraph_format.space_after = Pt(20)
        
        doc.add_heading("Professional Experience", level=1)
        doc.add_paragraph("- Worked as a backend dev at tech co.", style='List Bullet')
        doc.add_paragraph("- Developed database schemas and queries.", style='List Bullet')
        doc.add_paragraph("- Wrote python backend scripts.", style='List Bullet')
        
        doc.add_heading("Skills Summary", level=1)
        doc.add_paragraph("Python, FastAPI, SQL, Docker, HTML, JavaScript")
        
        doc.add_heading("Education", level=1)
        doc.add_paragraph("B.S. in Computer Science - State University (2018 - 2022)")
        
        doc.add_heading("Certifications", level=1)
        doc.add_paragraph("PostgreSQL Certified Professional")
        
        doc.save(path)
        logger.info(f"Created default base resume template at {path}")

