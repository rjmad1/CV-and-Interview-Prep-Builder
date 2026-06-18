import io
from pypdf import PdfReader
from docx import Document

def parse_pdf(file_bytes: bytes) -> str:
    """Extracts raw text from a PDF document using pypdf."""
    pdf_file = io.BytesIO(file_bytes)
    reader = PdfReader(pdf_file)
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n".join(text_parts)

def parse_docx(file_bytes: bytes) -> str:
    """Extracts text from a DOCX document using python-docx."""
    docx_file = io.BytesIO(file_bytes)
    doc = Document(docx_file)
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text_parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text_parts.append(cell.text)
    return "\n".join(text_parts)

def parse_document(file_bytes: bytes, filename: str) -> str:
    """Parses document based on its extension."""
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        return parse_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        return parse_docx(file_bytes)
    else:
        # Fallback to plain text
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1")
