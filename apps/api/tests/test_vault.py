import pytest
import uuid
import os
import tempfile
from fastapi.testclient import TestClient
from apps.api.src.main import app
from apps.api.src.database import get_db, Base, SessionLocal, engine
from apps.api.src.models import (
    User, Document, JobDescription, ResumeTemplate, ResumeVersion
)
from sqlalchemy import select
from apps.api.src.utils.pdf_generator import convert_docx_to_pdf, generate_cover_letter_pdf
from docx import Document as DocxDocument

@pytest.fixture(name="client", scope="module")
def fixture_client():
    # Make sure all tables are created in the test database
    Base.metadata.create_all(bind=engine)
    
    # Seed a developer user to satisfy get_current_user
    db = SessionLocal()
    dev_email = "developer@career-intelligence.studio"
    user = db.query(User).filter(User.email == dev_email).first()
    if not user:
        user = User(
            id=uuid.uuid4(),
            email=dev_email,
            hashed_password="pbkdf2:sha256:mock_hash_for_dev",
            first_name="Lead",
            last_name="Architect",
            role="admin"
        )
        db.add(user)
        db.commit()
    
    # Seed a default resume template
    template = db.query(ResumeTemplate).filter(ResumeTemplate.name == "Professional Default Template").first()
    if not template:
        template = ResumeTemplate(
            id=uuid.uuid4(),
            user_id=user.id,
            name="Professional Default Template",
            file_path="default_resume.docx",
            is_active=True
        )
        db.add(template)
        db.commit()
        
    db.close()
    
    # Mock background parsing pipeline instance method directly via __dict__ to bypass Pydantic's __setattr__ validation
    from unittest.mock import AsyncMock
    from apps.api.src.graph.document_processing import document_processing_graph
    document_processing_graph.__dict__['ainvoke'] = AsyncMock(return_value={"status": "completed"})
    
    client = TestClient(app)
    yield client
    
    # Clean up overrides
    app.dependency_overrides.clear()
    if 'ainvoke' in document_processing_graph.__dict__:
        del document_processing_graph.__dict__['ainvoke']

def test_cleanup_reset_endpoint(client):
    # 1. Trigger the reset endpoint
    response = client.post("/api/cleanup/reset")
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "success"
    assert "baseline state" in data["message"]

def test_strict_one_to_one_active_resume(client):
    # Enforce strict 1-to-1 active resume model: previous resumes must be archived when a new one is scanned
    # Ingest CV 1
    response1 = client.post(
        "/api/documents/scan",
        files={"file": ("resume_alpha.txt", b"Content Alpha", "text/plain")},
        params={"document_type": "resume"}
    )
    assert response1.status_code == 202
    
    # Ingest CV 2
    response2 = client.post(
        "/api/documents/scan",
        files={"file": ("resume_beta.txt", b"Content Beta", "text/plain")},
        params={"document_type": "resume"}
    )
    assert response2.status_code == 202
    
    # Fetch all documents to check archiving status
    db = SessionLocal()
    res = db.execute(select(Document).filter(Document.document_type == "resume").order_by(Document.created_at))
    documents = res.scalars().all()
    db.close()
    
    # We should have exactly 2 documents, and the first one (Alpha) should be archived, and the second one (Beta) should be active.
    assert len(documents) == 2
    assert documents[0].filename == "resume_alpha.txt"
    assert documents[0].is_archived is True
    assert documents[1].filename == "resume_beta.txt"
    assert documents[1].is_archived is False

def test_strict_one_to_one_active_jd(client):
    # Enforce strict 1-to-1 active JD model: previous JDs must be archived when a new one is analyzed
    # Analyze JD 1
    response1 = client.post(
        "/api/jd/analyze",
        json={"jd_text": "Required skills: Python, FastAPI."}
    )
    assert response1.status_code == 200
    
    # Analyze JD 2
    response2 = client.post(
        "/api/jd/analyze",
        json={"jd_text": "Required skills: Go, Kubernetes."}
    )
    assert response2.status_code == 200
    
    # Fetch all job descriptions to check archiving status
    db = SessionLocal()
    res = db.execute(select(JobDescription).order_by(JobDescription.created_at))
    jds = res.scalars().all()
    db.close()
    
    # We should have exactly 2 JDs, the first one archived, and the second one active
    assert len(jds) == 2
    assert jds[0].raw_text == "Required skills: Python, FastAPI."
    assert jds[0].is_archived is True
    assert jds[1].raw_text == "Required skills: Go, Kubernetes."
    assert jds[1].is_archived is False

def test_pdf_rendering_fidelity():
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "test.docx")
        pdf_path = os.path.join(tmpdir, "test.pdf")
        
        # Create a basic structured docx
        doc = DocxDocument()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("john@example.com | 123-456-7890")
        doc.add_paragraph("Professional Experience")
        doc.add_paragraph("- Built scalable backend APIs in python.")
        doc.save(docx_path)
        
        # Run conversion
        convert_docx_to_pdf(docx_path, pdf_path)
        
        assert os.path.exists(pdf_path)
        assert os.path.getsize(pdf_path) > 0
        
        # Run cover letter generation PDF
        cl_pdf_path = os.path.join(tmpdir, "cover_letter.pdf")
        generate_cover_letter_pdf("Dear Hiring Manager,\n\nI am writing to express my interest...\n\nSincerely,\nJohn", cl_pdf_path)
        assert os.path.exists(cl_pdf_path)
        assert os.path.getsize(cl_pdf_path) > 0
