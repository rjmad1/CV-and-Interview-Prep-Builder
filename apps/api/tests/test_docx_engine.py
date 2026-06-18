import os
import tempfile
import pytest
from docx import Document
from apps.api.src.engine.docx_engine import DocxEngine

def test_docx_engine_merge_changes():
    engine = DocxEngine()
    
    # Use temporary directory for testing templates and output files
    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, "test_template.docx")
        output_path = os.path.join(tmpdir, "test_output.docx")
        
        # 1. Create default template
        engine._create_default_template(template_path)
        assert os.path.exists(template_path)
        
        # Inspect template properties
        doc = Document(template_path)
        assert doc.paragraphs[2].text.strip() == "Professional Experience"
        # The first bullet:
        assert "backend dev" in doc.paragraphs[3].text
        # Remember style of first bullet run
        original_run = doc.paragraphs[3].runs[0]
        original_bold = original_run.bold
        original_font_name = original_run.font.name
        
        # 2. Merge changes
        optimized_sections = {
            "Professional Experience": (
                "- Engineered FastAPI microservices with RLS security\n"
                "- Led a team of 4 software developers to deploy ECS container workloads\n"
                "- Scaled Postgres db queries improving search metrics by 50%\n"
                "- Wrote comprehensive unit and integration test suites"
            ),
            "Skills Summary": "Python, FastAPI, Docker, SQL, AWS ECS, CI/CD pipelines"
        }
        
        engine.merge_changes(template_path, optimized_sections, output_path)
        assert os.path.exists(output_path)
        
        # 3. Verify modifications
        new_doc = Document(output_path)
        
        # Verify single-paragraph Skills Summary replacement
        skills_p = None
        for idx, p in enumerate(new_doc.paragraphs):
            if p.text.strip() == "Skills Summary":
                skills_p = new_doc.paragraphs[idx + 1]
                break
        assert skills_p is not None
        assert "AWS ECS" in skills_p.text
        
        # Verify list bullets replacement and formatting preservation
        bullets = []
        heading_found = False
        for p in new_doc.paragraphs:
            if p.text.strip() == "Professional Experience":
                heading_found = True
                continue
            if heading_found:
                if p.style.name.startswith("Heading") or p.text.strip() in ["Skills Summary", "Education", "Certifications"]:
                    break
                if p.text.strip() != "":
                    bullets.append(p)
                    
        # We optimized 4 bullets, template had 3. Verify it inserted the 4th bullet
        assert len(bullets) == 4
        assert "FastAPI microservices" in bullets[0].text
        assert "ECS container workloads" in bullets[1].text
        assert "comprehensive unit" in bullets[3].text
        
        # Check style preservation of first run on replaced bullet
        first_replaced_run = bullets[0].runs[0]
        assert first_replaced_run.bold == original_bold
        assert first_replaced_run.font.name == original_font_name
        
        # Verify bullet characters prefix were correctly lstriped
        assert not bullets[0].text.startswith("-")
        assert not bullets[0].text.startswith("*")

def test_docx_engine_xml_placeholders_and_regression():
    engine = DocxEngine()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, "test_template_xml.docx")
        output_path = os.path.join(tmpdir, "test_output_xml.docx")
        
        # Create default template
        engine._create_default_template(template_path)
        
        # Modify it to add placeholders
        doc = Document(template_path)
        doc.add_paragraph("Placeholder Test: {{candidate_name}}")
        doc.save(template_path)
        
        # Run merging with placeholders and sections
        optimized_sections = {
            "Skills Summary": "FastAPI, PyTest, Docker, SQL"
        }
        placeholders = {
            "candidate_name": "Antigravity Architect"
        }
        
        # Verify unzip_and_replace_xml succeeds
        success = engine.unzip_and_replace_xml(
            template_path, 
            optimized_sections, 
            output_path, 
            placeholders=placeholders
        )
        assert success is True
        assert os.path.exists(output_path)
        
        # Verify placeholder was replaced
        new_doc = Document(output_path)
        found = False
        for p in new_doc.paragraphs:
            if "Antigravity Architect" in p.text:
                found = True
                break
        assert found is True

        # Verify layout regression check detects anomalies
        original_xml = "<body xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"><w:tbl></w:tbl><w:pgMar w:top=\"100\"/></body>"
        # Mutate table (regression - missing table)
        mutated_xml_table = "<body xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"><w:pgMar w:top=\"100\"/></body>"
        assert engine.verify_layout_regression(original_xml, mutated_xml_table) is False
        
        # Mutate margins (regression)
        mutated_xml_margin = "<body xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"><w:tbl></w:tbl><w:pgMar w:top=\"200\"/></body>"
        assert engine.verify_layout_regression(original_xml, mutated_xml_margin) is False

